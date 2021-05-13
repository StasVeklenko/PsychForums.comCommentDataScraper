import sys, time, os#, openpyxl, collections
from docx import *
from docx.shared import RGBColor
from Projects.CustomClasses.Page import Page
from Projects.CustomClasses.Browser import Browser
from selenium.webdriver.common.keys import Keys
from selenium.common.exceptions import TimeoutException, ElementNotVisibleException, NoSuchElementException, NoSuchFrameException

links = ["http://www.psychforums.com/bipolar/"]
search_terms = ["Omega 3", "ketogenic", "vegetarian"]

browser = Browser(starting_page=links[0])

def word_input(document, doc_name, title, content, author, quote):
    author_para = document.add_paragraph(text=author)
    title_para = document.add_paragraph(text=title)
    document.add_paragraph(text="----------------------------------------------------------------------")
    if quote:
        quote_para = document.add_paragraph(text=quote)
        quote_para.style = document.styles['List Paragraph']
        document.add_paragraph()
    content_para = document.add_paragraph(text=content)   
    author_para.style = document.styles['Heading 2']
    author_para.style.font.underline = True
    author_para.style.font.color.rgb = RGBColor(255, 160, 122)
    title_para.style = document.styles['Heading 3']
    content_para.style = document.styles['Normal']
    document.add_paragraph(text="======================================================================") 
    document.save(doc_name)

forum = Page(browser)
keyword = "vegetarian"
if keyword == "vegetarian":
    document = Document()
    search_box = forum.find_elements_on_page(False, "xpath", "//form[@id='forum-search']/fieldset/input[contains(@class, 'inputbox')]") #finding a 'search' button
    search_box.send_keys(Keys.CONTROL + "a")
    search_box.send_keys(Keys.DELETE)
    search_box.send_keys(keyword)
    forum.find_elements_on_page(False, "xpath", "//form[@id='forum-search']/fieldset/input[contains(@class, 'button2')]").click()  #clicking on a 'search' button
    
    if forum.explicit_wait("//div[contains(@class,'search post')]", "xpath"):  #if posts started showing up
        pages_elements = forum.find_elements_on_page(False, "css", "div.topic-actions div:nth-child(2) span a:last-of-type") #returns last page number or an empty list
        pages = 1

        if pages_elements:
            pages = int(pages_elements.text)  #getting the number of pages in search results

        for page in range(1, pages+1):
            for post_shortcut in forum.find_elements_on_page(True, "xpath", "//div[contains(@class,'search post')]"): #for each post that has been found
                #checking if at least 1 keyword is found in a post
                if keyword == "Omega 3":
                    cond = any(key in post_shortcut.find_element_by_xpath(".//div[@class='postbody']").text for key in ["Omega  3", "omega  3", "Omega 3","omega 3","Omega-3", "omega-3", "Omega -3", "Omega- 3", "Omega - 3", "omega -3", "omega- 3", "omega - 3", "Omega3", "omega3", "Omega Three", "Omega three", "omega Three", "omega three", "Omega-Three", "Omega-three", "omega-Three", "omega-three"])
                elif keyword == "ketogenic":
                    cond = any(key in post_shortcut.find_element_by_xpath(".//div[@class='postbody']").text for key in ["Ketogenic", "ketogenic"])
                elif keyword == "vegetarian":
                    cond = any(key in post_shortcut.find_element_by_xpath(".//div[@class='postbody']").text for key in ["Vegetarian", "vegetarian"])
                
                if cond: #i.e. if any of the above keywords is in post
                    date = post_shortcut.find_element_by_xpath(".//div[@class='inner']/dl[@class='postprofile']/dd[1]").text #date of post
                    link = post_shortcut.find_element_by_xpath(".//ul[@class='searchresults']/li/a").get_attribute("href") #link to full post
                    browser.open_new_tab(tab_link=link, switch_to_new_tab=True)
                    post_page = Page(browser)

                    for tag in post_page.find_elements_on_page(multiple=True, _method="css", elements_pointer="p.author"):
                        if date in tag.text:
                            try:
                                post_quote = tag.find_element_by_xpath(".//following-sibling::div/blockquote/div").text  #text the post is quoting
                                whole_text = tag.find_element_by_xpath(".//following-sibling::div").text #whole text of post (quote + response)
                                post_text = whole_text[len(post_quote):] #extract text of quote from whole_text
                            except NoSuchElementException:
                                #print("NoSuchElementException was raised on " + date)
                                post_quote = None
                                post_text = tag.find_element_by_xpath(".//following-sibling::div").text
                            post_author = tag.find_element_by_xpath("./strong/a").text
                            post_title = tag.find_element_by_xpath(".//preceding-sibling::h3/a").text
                            word_input(document, 'Psychforums '+keyword+'.docx', post_title, post_text, post_author, post_quote)
                            break
                   
                    browser.close_tab("www.psychforums.com_2", switch_to_tab_key="www.psychforums.com_1")
            if (page != pages):
                forum.find_elements_on_page(False, "xpath", "//div[@class='topic-actions']/div[contains(@class,'pagination')]/span/a[contains(text(), '"+str(page+1)+"')]").click() #click to go to next page 
            time.sleep(5)
        browser.chrome.quit()
    else:
        print("element was not found")

